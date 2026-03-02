from __future__ import annotations

from io import BytesIO

from docx import Document

from app.models import Job, JobChannel, TranscriptWord


def _fmt_srt_time(seconds: float) -> str:
    total_ms = max(0, int(round(seconds * 1000)))
    h = total_ms // 3_600_000
    m = (total_ms % 3_600_000) // 60_000
    s = (total_ms % 60_000) // 1000
    ms = total_ms % 1000
    return f"{h:02d}:{m:02d}:{s:02d},{ms:03d}"


def _iter_channel_words(channel: JobChannel) -> list[TranscriptWord]:
    return sorted(channel.words, key=lambda w: w.seq)


def render_txt(job: Job) -> bytes:
    lines: list[str] = []
    lines.append(f"Job: {job.id}")
    lines.append(f"File: {job.source_filename}")
    lines.append("")

    for channel in sorted(job.channels, key=lambda c: c.channel_index):
        lines.append(f"[Channel {channel.channel_index + 1}]")
        words = _iter_channel_words(channel)
        if words:
            grouped: dict[str, list[str]] = {}
            order: list[str] = []
            for w in words:
                spk = w.speaker_label or "spk_unknown"
                if spk not in grouped:
                    grouped[spk] = []
                    order.append(spk)
                grouped[spk].append(w.text)
            for spk in order:
                lines.append(f"{spk}: {' '.join(grouped[spk]).strip()}")
        else:
            lines.append((channel.transcript_text or "").strip())
        lines.append("")

    return "\n".join(lines).encode("utf-8")


def _build_srt_entries(channel: JobChannel) -> list[tuple[float, float, str]]:
    words = _iter_channel_words(channel)
    if not words:
        return []

    entries: list[tuple[float, float, str]] = []
    bucket: list[TranscriptWord] = []

    def flush_bucket() -> None:
        if not bucket:
            return
        start = bucket[0].start_sec
        end = bucket[-1].end_sec
        speaker = bucket[0].speaker_label or f"spk_ch_{channel.channel_index + 1}"
        text = " ".join(w.text for w in bucket).strip()
        entries.append((start, end, f"{speaker}: {text}"))
        bucket.clear()

    for w in words:
        if not bucket:
            bucket.append(w)
            continue

        prev = bucket[-1]
        gap = w.start_sec - prev.end_sec
        speaker_changed = (w.speaker_label or "") != (prev.speaker_label or "")
        if len(bucket) >= 10 or gap > 0.9 or speaker_changed:
            flush_bucket()
        bucket.append(w)

    flush_bucket()
    return entries


def render_srt(job: Job) -> bytes:
    all_entries: list[tuple[float, float, str]] = []
    fallback_cursor = 0.0

    for channel in sorted(job.channels, key=lambda c: c.channel_index):
        entries = _build_srt_entries(channel)
        if entries:
            all_entries.extend(entries)
            continue

        fallback_text = (channel.transcript_text or "").strip()
        if fallback_text:
            start = fallback_cursor
            end = start + max(2.0, min(8.0, float(job.source_duration_sec)))
            speaker = f"spk_ch_{channel.channel_index + 1}"
            all_entries.append((start, end, f"{speaker}: {fallback_text}"))
            fallback_cursor = end + 0.1

    if not all_entries:
        all_entries.append((0.0, 2.0, 'spk_unknown: [no transcript text available]'))

    all_entries.sort(key=lambda x: (x[0], x[1]))
    lines: list[str] = []
    for idx, (start, end, text) in enumerate(all_entries, start=1):
        lines.append(str(idx))
        lines.append(f"{_fmt_srt_time(start)} --> {_fmt_srt_time(end)}")
        lines.append(text)
        lines.append("")

    return "\n".join(lines).encode("utf-8")


def render_docx(job: Job) -> bytes:
    doc = Document()
    doc.add_heading("Hebrew Transcription", level=1)
    doc.add_paragraph(f"Job ID: {job.id}")
    doc.add_paragraph(f"Source file: {job.source_filename}")

    for channel in sorted(job.channels, key=lambda c: c.channel_index):
        doc.add_heading(f"Channel {channel.channel_index + 1}", level=2)
        words = _iter_channel_words(channel)
        if words:
            by_speaker: dict[str, list[str]] = {}
            order: list[str] = []
            for w in words:
                spk = w.speaker_label or "spk_unknown"
                if spk not in by_speaker:
                    by_speaker[spk] = []
                    order.append(spk)
                by_speaker[spk].append(w.text)
            for spk in order:
                doc.add_paragraph(f"{spk}: {' '.join(by_speaker[spk]).strip()}")
        else:
            doc.add_paragraph((channel.transcript_text or "").strip())

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
